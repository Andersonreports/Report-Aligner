import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

def set_style(doc, name, font_name, size, bold, color, alignment, space_after=None):
    style = doc.styles[name]
    style.font.name = font_name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(*color)
    style.paragraph_format.alignment = alignment
    if space_after is not None:
        style.paragraph_format.space_after = Pt(space_after)

# ---- Template ----
tpl = docx.Document()
set_style(tpl, 'Normal', 'Calibri', 11, False, (0, 0, 0), WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
set_style(tpl, 'Heading 1', 'Calibri', 18, True, (0x1F, 0x4E, 0x79), WD_ALIGN_PARAGRAPH.LEFT, space_after=12)
set_style(tpl, 'Heading 2', 'Calibri', 14, True, (0x2E, 0x74, 0xB5), WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

section = tpl.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)
section.orientation = WD_ORIENT.PORTRAIT

tpl.add_paragraph('Report Title', style='Heading 1')
tpl.add_paragraph('Section', style='Heading 2')
p = tpl.add_paragraph('Body text goes here in the template.')
table = tpl.add_table(rows=2, cols=2)
table.style = 'Light Grid Accent 1'
tpl.save('template.docx')

# ---- Report (misaligned, to be fixed) ----
rep = docx.Document()
section = rep.sections[0]
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

h1 = rep.add_paragraph('Report Title', style='Heading 1')
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in h1.runs:
    r.font.size = Pt(24)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0, 0, 0)

h2 = rep.add_paragraph('Section', style='Heading 2')
h2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

body = rep.add_paragraph('Body text goes here in the report and is left aligned instead of justified, in the wrong font.')
body.alignment = WD_ALIGN_PARAGRAPH.LEFT
for r in body.runs:
    r.font.name = 'Arial'
    r.font.size = Pt(13)

# untagged heading (Normal style, bold, short) for heuristic test
untagged = rep.add_paragraph('Untagged Heading')
for r in untagged.runs:
    r.font.bold = True
    r.font.size = Pt(16)

table = rep.add_table(rows=2, cols=2)
table.style = 'Table Grid'

rep.save('report.docx')
print("fixtures written")

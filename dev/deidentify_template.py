"""
Replace real patient-identifying / case-specific text in a report docx with
generic placeholders, preserving formatting as much as practical, so the
file stays usable as a formatting template.

Operates on paragraph text at the XML level (all <w:t> descendants of a
paragraph, via a `.//` search) rather than python-docx's `paragraph.runs`,
because `.runs` only returns direct-child runs and silently skips runs
wrapped in <w:hyperlink> (e.g. "OMIM#..." fields are usually hyperlinks).
When a paragraph's full text matches a replacement, all its text nodes are
collapsed into the first one to avoid re-splitting text across runs with
different original formatting.

Usage: python deidentify_template.py <input.docx> <output.docx>
"""
import sys
import docx
from docx.oxml.ns import qn

REPLACEMENTS = [
    ("Ms. Kerina", "Ms. Sample Patient"),
    ("Kerina", "Sample Patient"),
    ("AVP0000008750", "AVP0000000000"),
    ("Sri Baalaji Clinical Laboratory - Villupuram", "Sample Clinic - City"),
    ("260604571", "000000000"),
    ("22/05/2026", "01/01/2026"),
    ("23/05/2026", "02/01/2026"),
    ("16/06/2026", "10/01/2026"),

    ("hereditary neuropathy. Her nerve biopsy suggestive of chronic demyelinating neuropathy with secondary axonal loss.",
     "a suspected hereditary condition. Clinical findings were consistent with the reported phenotype."),
    ("This variant has been previously reported in patient affected with Charcot-Marie-Tooth Disease Type 4C in homozygous state.",
     "This variant has been previously reported in a patient affected with the same disorder in homozygous state."),

    ("SH3TC2", "GENE1"),
    ("c.2710C>T", "c.100A>T"),
    ("p.Arg904Ter", "p.Lys34Ter"),
    ("c.3287_3308del", "c.200_221del"),
    ("p.Asn1096MetfsTer13", "p.Gly67MetfsTer5"),
    ("Charcot-Marie-Tooth disease, type 4C (CMT4C)", "Sample Autosomal Recessive Disorder (SARD)"),
    ("Charcot-Marie-Tooth disease type 4C (CMT4C)", "Sample Autosomal Recessive Disorder (SARD)"),
    ("Charcot-Marie-Tooth disease, type 4C", "Sample Autosomal Recessive Disorder"),
    ("Charcot-Marie-Tooth disease type 4C", "Sample Autosomal Recessive Disorder"),
    ("Charcot-Marie-Tooth", "Sample Autosomal Recessive Disorder"),
    ("OMIM#601596", "OMIM#000000"),
    ("OMIM*608206", "OMIM*000000"),
    ("Chr5:NC_000005.10:g.149027022G>A", "Chr1:NC_000001.11:g.100000A>T"),
    ("Chr5:NC_000005.10:g.149010289del(22bp)", "Chr1:NC_000001.11:g.200000del(22bp)"),
    ("NM_024577.4", "NM_000000.1"),
    ("ClinVar: 21696", "ClinVar: 0000000"),
    ("PubMed: 29515423", "PubMed: 00000000"),
    ("Ref(G): 71, Alt(A): 43, VAF: 37.72%", "Ref(G): 50, Alt(A): 50, VAF: 50.00%"),
    ("Ref(TGATGCCTGTGGCGGGTCCCAT): 73,", "Ref(N): 50,"),
    ("Alt(-): 42, VAF: 36.52%", "Alt(-): 50, VAF: 50.00%"),
    ("12 Gb", "10 Gb"),
    ("96.79%", "95.00%"),
]


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        yield from iter_table_paragraphs(t)
    for s in doc.sections:
        for part in (s.header, s.footer):
            for p in part.paragraphs:
                yield p
            for t in part.tables:
                yield from iter_table_paragraphs(t)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def paragraph_text_nodes(paragraph):
    return paragraph._p.findall('.//' + qn('w:t'))


def apply_replacements(text):
    for old, new in REPLACEMENTS:
        if old in text:
            text = text.replace(old, new)
    return text


def main():
    src, dst = sys.argv[1], sys.argv[2]
    doc = docx.Document(src)

    for p in iter_all_paragraphs(doc):
        nodes = paragraph_text_nodes(p)
        if not nodes:
            continue
        original = ''.join(n.text or '' for n in nodes)
        replaced = apply_replacements(original)
        if replaced == original:
            continue
        nodes[0].text = replaced
        for n in nodes[1:]:
            n.text = ''

    doc.save(dst)
    print("wrote", dst)


if __name__ == "__main__":
    main()
